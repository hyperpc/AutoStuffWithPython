import sys, os, docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generateInvitation():
    currentPath = os.path.dirname(sys.argv[0])
    templatePath = os.path.join(currentPath, 'invitation_template.docx')
    invitationPath = os.path.join(currentPath, 'invitation_generated.docx')
    doc = docx.Document(templatePath)
    guestsFilePath = os.path.join(currentPath, 'guests.txt')
    guestsFileObj = open(guestsFilePath, 'r')
    names = guestsFileObj.readlines()
    for name in names:
        if name == None:
            continue
        if name.strip() == '':
            continue
        # top
        para_top = doc.add_paragraph('It would be a pleasure to have the company of')
        para_top.style = 'invitation_title'
        para_top_format = para_top.paragraph_format
        para_top_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # name
        para_name = doc.add_paragraph(name.strip())
        para_name.style = 'invitation_guest_name'
        para_name_format = para_name.paragraph_format
        para_name_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # mid
        para_mid = doc.add_paragraph('')
        para_mid_run_head = para_mid.add_run('at')
        #para_mid_run_head.style = 'QuoteChar'  # english version
        #para_mid_run_head.style = '智能链接'  # chinese version
        para_mid_run_head.underline = True
        para_mid_run_tail = para_mid.add_run(' 11010 Memory Lane on the Evening of')
        para_mid.style = 'invitation_title'
        para_mid_format = para_mid.paragraph_format
        para_mid_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # date
        para_date = doc.add_paragraph('April 1st')
        para_date.style = 'invitation_date'
        para_date_format = para_date.paragraph_format
        para_date_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # time
        para_time = doc.add_paragraph('')
        para_time_run_head = para_time.add_run('at')
        #para_time_run_head.style = 'QuoteChar'  # english version
        #para_time_run_head.style = '智能链接'  # chinese version
        para_time_run_head.underline = True
        para_time_run_tail = para_time.add_run(' 7 o\'clock')
        para_time.style = 'invitation_title'
        para_time_format = para_time.paragraph_format
        para_time_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
    doc.save(invitationPath)
    print('Done~')

def main():
    generateInvitation()

if __name__ == '__main__':
    main()