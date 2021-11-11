import docx, sys, os

def addNewPage():
    currentPath = os.path.dirname(sys.argv[0])
    twopageFilePath = os.path.join(currentPath, 'twoPage_my.docx')
    doc = docx.Document()
    doc.add_paragraph('This is on the first paragraph!')
    #doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
    doc.add_page_break()
    doc.add_paragraph('This is on the second paragraph!')
    doc.save(twopageFilePath)

def main():
    addNewPage()

if __name__ == '__main__':
    main()
