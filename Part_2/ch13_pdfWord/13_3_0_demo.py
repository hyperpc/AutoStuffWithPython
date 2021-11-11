import docx, sys, os

def readDocx():
    currentPath = os.path.dirname(sys.argv[0])
    # 13.3.1
    doc = docx.Document(os.path.join(currentPath, 'demo.docx'))
    print('len(doc.paragraphs) = ', len(doc.paragraphs))
    print('doc.paragraphs[0].text = ', doc.paragraphs[0].text)
    print('doc.paragraphs[1].text = ', doc.paragraphs[1].text)
    print('len(doc.paragraphs[1].runs) = ', len(doc.paragraphs[1].runs))
    print('doc.paragraphs[1].runs[0].text = ', doc.paragraphs[1].runs[0].text)
    print('doc.paragraphs[1].runs[1].text = ', doc.paragraphs[1].runs[1].text)
    print('doc.paragraphs[1].runs[2].text = ', doc.paragraphs[1].runs[2].text)
    print('doc.paragraphs[1].runs[3].text = ', doc.paragraphs[1].runs[3].text)
    print('doc.paragraphs[1].runs[4].text = ', doc.paragraphs[1].runs[4].text)
    #13.3.2
    fullText = []
    for para in doc.paragraphs:
        fullText.append('  '+para.text)
    print('\n'.join(fullText))
    #13.3.5
    print('(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text, doc.paragraphs[1].runs[4].text) = ', (doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text, doc.paragraphs[1].runs[4].text))
    doc.paragraphs[1].runs[0].style = 'QuoteChar'
    doc.paragraphs[1].runs[1].underline = True
    doc.paragraphs[1].runs[3].underline = True
    restyleFilePath = os.path.join(currentPath, 'restyled_my.docx')
    doc.save(restyleFilePath)
    print('restyleFilePath = ', restyleFilePath)

def main():
    readDocx()

if __name__ == '__main__':
    main()
