import docx, sys, os

def addHeadings():
    currentPath = os.path.dirname(sys.argv[0])
    headingFilePath = os.path.join(currentPath, 'headings_my.docx')
    doc = docx.Document()
    doc.add_heading('Header 0', 0)
    doc.add_heading('Header 1', 1)
    doc.add_heading('Header 2', 2)
    doc.add_heading('Header 3', 3)
    doc.add_heading('Header 4', 4)
    doc.save(headingFilePath)

def main():
    addHeadings()

if __name__ == '__main__':
    main()
